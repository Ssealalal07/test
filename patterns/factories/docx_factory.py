from docx import Document
from docx.shared import Inches
from factories.abstract_factory import Header, Table, Image, ReportFactory


class DocxHeader(Header):
    def __init__(self, doc):
        self.doc = doc

    def render(self, text):
        self.doc.add_heading(text, level=1)

class DocxTable(Table):
    def __init__(self, doc):
        self.doc = doc

    def render(self, data):
        table = self.doc.add_table(rows=1, cols=len(data[0]))
        hdr_cells = table.rows[0].cells
        # hdr_cells список ячеек
        for i, item in enumerate(data[0]):
            # ['id', 'name', 'city']
            hdr_cells[i].text = str(item)

        for row in data[1:]:
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

class DocxImage(Image):
    def __init__(self, doc):
        self.doc = doc

    def render(self, path):
        self.doc.add_picture(path, width=Inches(3))

class DocxReportFactory(ReportFactory):
    def __init__(self):
        self.doc = Document()

    def create_header(self) -> Header:
        return DocxHeader(self.doc)

    def create_table(self) -> Table:
        return DocxTable(self.doc)

    def create_image(self) -> Image:
        return DocxImage(self.doc)

    def save(self, filename):
        self.doc.save(filename)

